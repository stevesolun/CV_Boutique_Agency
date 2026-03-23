from __future__ import annotations
from pathlib import Path
from typing import Dict, Any
from docx import Document
from docx.shared import Pt

def export_resume_to_docx(resume_data: Dict[str, Any], output_path: str) -> str:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph()
    r = p.add_run(resume_data.get("name", ""))
    r.bold = True
    r.font.size = Pt(16)

    for line in resume_data.get("header_lines", []):
        doc.add_paragraph(line)

    for section in resume_data.get("sections", []):
        doc.add_heading(section.get("title", ""), level=2)
        for item in section.get("items", []):
            if isinstance(item, str):
                doc.add_paragraph(item, style="List Bullet")
            elif isinstance(item, dict):
                if "company" in item or "title" in item or "dates" in item:
                    # Structured experience item: render company | title on one bold line, dates alongside
                    name_parts = [part for part in [item.get("company"), item.get("title")] if part]
                    heading_text = " | ".join(name_parts) if name_parts else ""
                    dates_text = item.get("dates", "")
                    p = doc.add_paragraph()
                    rr = p.add_run(heading_text)
                    rr.bold = True
                    if dates_text:
                        p.add_run(f"  {dates_text}")
                elif "heading" in item:
                    p = doc.add_paragraph()
                    rr = p.add_run(item["heading"])
                    rr.bold = True
                for bullet in item.get("bullets", []):
                    doc.add_paragraph(bullet, style="List Bullet")

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return str(out)
